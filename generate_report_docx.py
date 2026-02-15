"""
agents_homepage 통합회고보고서 — 모듈홈페이지 보고서와 동일한 형식
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(60, 36, 21)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ── 유틸리티 함수 ──
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'F5E6D0'
        })
        shading.append(shading_elm)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = cell_data
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


# ═══════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('알앤디파크 주식회사')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(139, 109, 82)

doc.add_paragraph()

main_title = doc.add_paragraph()
main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main_title.add_run('디저트 쇼핑몰 홈페이지 제작 프로젝트\n통합 회고 보고서')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(60, 36, 21)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('프로젝트 개요  |  병목 현상  |  2차 프로젝트 비교  |  교훈')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(139, 109, 82)

for _ in range(4):
    doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2026년 2월 15일')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

info1 = doc.add_paragraph()
info1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info1.add_run('프로젝트: agents_homepage (Next.js + Vercel)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info2.add_run('배포 URL: Vercel 자동 배포 (GitHub 연동)')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 1. 프로젝트 개요
# ═══════════════════════════════════════════════════════

doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_paragraph(
    '알앤디파크의 세 번째 웹 프로젝트로, 프리미엄 디저트 판매 쇼핑몰 "Douceur"입니다. '
    '1차 프로젝트(rndpark-homepage)와 2차 프로젝트(module_homepage)의 교훈을 바탕으로, '
    '기술 스택을 Next.js 14 + Tailwind CSS로 전환하고 개발 도구를 Claude Code(에이전트)로 '
    '업그레이드하여 진행했습니다.'
)

add_table(doc,
    ['항목', '내용'],
    [
        ['프로젝트명', 'agents_homepage (디저트 쇼핑몰 "Douceur")'],
        ['플랫폼', 'Next.js 14 (App Router) + React 18 + Tailwind CSS 3.4 + Vercel'],
        ['개발 방식', 'AI 에이전트(Claude Code) 활용 — 파일 읽기/쓰기/Git 자동화'],
        ['디자인 컨셉', '"Patisserie Moderne" — 파리 파티스리 우아함 + 현대적 미니멀리즘'],
        ['전략', '커스텀 테마 설계 → UI 컴포넌트 → 페이지 조립 → 전역 상태 → 마무리'],
        ['배포', 'Vercel 자동 배포 (git push → 빌드 → 배포 자동 완료)'],
    ]
)

# 완성된 페이지
doc.add_heading('완성된 페이지 (13개)', level=2)

add_table(doc,
    ['구분', '페이지명', '주요 기능'],
    [
        ['핵심', '메인 홈 (/)', '히어로 배너 슬라이드, 인기 상품, 신상품, 브랜드 스토리, 고객 후기'],
        ['핵심', '상품 목록 (/products)', '카테고리 필터, 정렬, 실시간 검색, URL 파라미터 연동'],
        ['핵심', '상품 상세 (/products/[id])', '이미지 갤러리, 옵션 선택, 수량, 장바구니/찜 추가, 리뷰'],
        ['주문', '장바구니 (/cart)', '상품 목록, 수량 변경, 삭제, 주문 요약'],
        ['주문', '주문/결제 (/checkout)', '배송 정보 입력, 결제 수단 선택, 주문 요약'],
        ['주문', '주문 완료 (/checkout/complete)', '주문 성공 확인, 주문번호, 상세 요약'],
        ['인증', '로그인 (/auth/login)', '이메일/비밀번호 입력, 유효성 검증'],
        ['인증', '회원가입 (/auth/signup)', '이름, 이메일, 비밀번호 입력, 검증'],
        ['회원', '마이페이지 (/mypage)', '주문 내역, 찜 목록, 회원정보 수정 탭'],
        ['정보', '브랜드 소개 (/about)', '브랜드 스토리, 철학, 팀 소개, 아틀리에'],
        ['정보', '공지사항/FAQ (/notice)', '공지사항·FAQ 탭 전환, 아코디언 토글'],
        ['정보', '고객센터 (/contact)', '문의 폼, 연락처, 운영 시간, 지도'],
        ['관리', '관리자 대시보드 (/admin)', '매출 현황, 주문 관리, 상품 관리'],
    ]
)

# 완성된 컴포넌트
doc.add_heading('완성된 컴포넌트 (23개)', level=2)

add_table(doc,
    ['카테고리', '컴포넌트명', '역할'],
    [
        ['UI (6개)', 'Button', 'Primary, Secondary, Ghost 변형 지원'],
        ['UI', 'Card', '범용 카드 컴포넌트, 호버 효과'],
        ['UI', 'Input', '폼 입력 필드, label 연결, 에러 표시'],
        ['UI', 'Badge', '상품 태그 (BEST, NEW, SALE)'],
        ['UI', 'Modal', 'ESC 닫기, 배경 클릭 닫기, 애니메이션'],
        ['UI', 'Toast', '알림 메시지, 자동 사라짐, 타입별 아이콘'],
        ['레이아웃 (3개)', 'Header', '반응형 네비게이션, 장바구니 뱃지, 다크모드 토글'],
        ['레이아웃', 'Footer', '브랜드 정보, 쇼핑 링크, 고객 서비스'],
        ['레이아웃', 'MobileMenu', '모바일 햄버거 메뉴, 접근성 지원'],
        ['홈 전용 (5개)', 'HeroBanner', '자동 전환 슬라이드, 인디케이터'],
        ['홈 전용', 'PopularProducts', '리뷰 기준 상위 4개, 가로 스크롤'],
        ['홈 전용', 'NewArrivals', '신상품 필터링, 4열 그리드'],
        ['홈 전용', 'BrandStory', '브랜드 소개, 통계 카드'],
        ['홈 전용', 'ReviewPreview', '고객 후기 카드 미리보기'],
        ['상품 (5개)', 'ProductCard', '배지, 할인율, 별점, 찜 버튼, onError 폴백'],
        ['상품', 'ProductGrid', '반응형 그리드 (4열→2열), 빈 상태 처리'],
        ['상품', 'ProductFilter', '카테고리 탭, 정렬, 실시간 검색'],
        ['상품', 'ProductGallery', '이미지 갤러리, 썸네일 선택'],
        ['상품', 'ProductReview', '리뷰 섹션, 별점 표시'],
        ['장바구니 (2개)', 'CartItem', '수량 변경, 삭제 버튼'],
        ['장바구니', 'CartSummary', '총금액 계산, 결제 버튼'],
        ['전역 (1개)', 'Providers', 'Theme, Auth, Cart, Wishlist Context 래퍼'],
        ['훅 (1개)', 'useScrollAnimation', 'Intersection Observer 스크롤 애니메이션'],
    ]
)

# 전역 상태 연동
doc.add_heading('전역 상태 연동 (4개 Context)', level=2)
doc.add_paragraph('React Context API를 사용하여 전역 상태를 관리하며, 모든 Context는 localStorage와 자동 동기화됩니다.')

doc.add_paragraph('CartContext: 장바구니 추가/삭제/수량변경/비우기, 총금액 자동 계산 → localStorage(douceur_cart)', style='List Bullet')
doc.add_paragraph('AuthContext: 로그인/회원가입/로그아웃/프로필수정 → localStorage(douceur_user, douceur_members)', style='List Bullet')
doc.add_paragraph('WishlistContext: 찜 토글/목록관리 → localStorage(douceur_wishlist)', style='List Bullet')
doc.add_paragraph('ThemeContext: 다크/라이트 모드 전환 → localStorage(douceur_theme) + document.documentElement 토글', style='List Bullet')

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 2. 병목 현상 분석
# ═══════════════════════════════════════════════════════

doc.add_heading('2. 병목 현상 분석', level=1)

doc.add_paragraph(
    '개발 중 시간을 크게 소모한 병목 현상 4건을 영향도 순으로 정리합니다. '
    '2차 프로젝트(module_homepage)에서 6건이 발생한 것에 비해, 기술 스택 전환과 교훈 적용 덕분에 '
    '병목 수 자체가 크게 줄었습니다.'
)

add_table(doc,
    ['#', '병목', '원인 → 해결', '영향도'],
    [
        ['1', 'URL 파라미터 ↔ useState 동기화 버그',
         '상품 목록에서 카테고리·정렬을 useState와 URL searchParams 두 곳에서 관리. Header 카테고리 링크 클릭 시 필터 미작동. searchParams를 단일 소스로 사용하고 useState 중복 제거로 해결',
         '★★★★☆'],
        ['2', 'useSearchParams Suspense 에러',
         'Next.js 14 App Router에서 useSearchParams는 반드시 Suspense 바운더리 안에서 사용해야 함. 빌드 시 에러 발생. Suspense 래퍼 컴포넌트 추가로 해결',
         '★★★☆☆'],
        ['3', 'next/font 빌드 환경 호환성',
         'next/font/google로 Playfair Display, Noto Sans KR, DM Sans 3종을 로드할 때 빌드 환경에서 에러 발생. CSS @import 방식으로 전환하여 해결',
         '★★★☆☆'],
        ['4', '이미지 파일 부재 시 깨짐',
         '개발 초기에 실제 이미지 파일 없이 진행. next/image에 onError 폴백 구조를 처음부터 설계하지 않아 이미지 깨짐. onError + placehold.co 폴백 패턴을 전 컴포넌트에 적용',
         '★★☆☆☆'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 3. 2차(module_homepage) vs 3차(agents_homepage) 비교 분석
# ═══════════════════════════════════════════════════════

doc.add_heading('3. 2차(module_homepage) vs 3차(agents_homepage) 비교 분석', level=1)

doc.add_paragraph(
    '2차 프로젝트의 교훈이 3차에서 어떻게 적용되었는지, '
    '그리고 새롭게 발생한 문제는 무엇인지 비교합니다.'
)

# 예방에 성공한 것
doc.add_heading('예방에 성공한 것 (6건)', level=2)

add_table(doc,
    ['항목', 'module_homepage (발생)', 'agents_homepage (예방)'],
    [
        ['에뮬레이터 ↔ 프로덕션 차이',
         'Firestore 복합 인덱스 없이 에뮬레이터 통과 → 배포 후 에러',
         'Firebase 미사용. localStorage 기반 → 문제 자체 소멸'],
        ['배포 설정 반복 실패',
         '인증 만료 + firebase.json site 누락 → 3~4회 실패',
         'Vercel 자동 배포 → git push만으로 완료. 설정 실패 0건'],
        ['Firebase Auth 재인증',
         '시간 경과 후 requires-recent-login 에러',
         '자체 AuthContext + localStorage → 외부 인증 의존 없음'],
        ['CSS 하드코딩 색상',
         '13개 모듈에 #03C75A 등 하드코딩 → 일괄 var() 변환',
         'tailwind.config.js에 커스텀 컬러 사전 정의 → 하드코딩 0건'],
        ['모듈 간 연동 충돌',
         '독립 개발 후 연동 시 의존성 충돌 → 백업 후 점진적 연동',
         'Context API로 처음부터 전역 연결 → 연동 충돌 0건'],
        ['변수 중복 선언 (1차 교훈)',
         '1차에서 예방 성공, 2차에서도 유지',
         '프레임워크 모듈 시스템 → 구조적으로 불가능'],
    ]
)

# 여전히 반복된 실수
doc.add_heading('여전히 반복된 실수 (2건, 단 영향도 크게 감소)', level=2)

add_table(doc,
    ['실수', 'module_homepage', 'agents_homepage', '변화'],
    [
        ['CSS/스타일 사후 수정',
         '하드코딩 색상 → 13개 모듈 일괄 전환 (수시간)',
         '다크모드 dark: 클래스를 전체 컴포넌트에 사후 추가',
         '범위 축소: 13개 모듈 → Tailwind 클래스 추가로 간소화'],
        ['컨텍스트 소진',
         '13개 모듈 작업량 → 컨텍스트 초과',
         'Claude Code 도구 기반으로 크게 개선. 단, 대규모 작업 시 여전히 한계',
         '빈도 감소: 매번 → 간헐적'],
    ]
)

# 시간 소모 비교
doc.add_heading('시간 소모 비교', level=2)

add_table(doc,
    ['rndpark (1차)', 'module_homepage (2차)', 'agents_homepage (3차)'],
    [
        ['규칙 없이 시작\n→ 8~13시간 낭비',
         '규칙 세팅 후 시작\n→ 4~8시간으로 감소',
         '규칙 + 프레임워크 + 에이전트\n→ 1~2시간으로 감소'],
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 4. 이번 프로젝트 고유 교훈
# ═══════════════════════════════════════════════════════

doc.add_heading('4. 이번 프로젝트 고유 교훈', level=1)

doc.add_paragraph(
    'module_homepage에서는 경험하지 못했던, agents_homepage에서 새로 얻은 교훈입니다.'
)

# 4-1. 개발 도구 전환의 효과
doc.add_heading('4-1. 개발 도구 전환의 효과 (채팅 → 에이전트)', level=2)

add_table(doc,
    ['장점', '단점'],
    [
        ['파일 읽기/쓰기를 에이전트가 직접 수행 → 복붙 실수 0건', '코드 변경 과정을 실시간으로 면밀히 추적하기 어려움'],
        ['Git 커밋을 기능 단위로 자동 관리 → 61개 구조화된 이력', '에이전트가 여러 파일을 동시에 수정할 때 의도 파악 필요'],
        ['에러 발생 시 자동 분석 + 수정 시도 → 디버깅 시간 단축', '에이전트의 수정 방향이 의도와 다를 수 있어 확인 필요'],
        ['배포가 git push 한 번으로 완료 → 배포 병목 완전 해소', '—'],
    ]
)

# 4-2. URL 파라미터 단일 소스 원칙
doc.add_heading('4-2. URL 파라미터 단일 소스 원칙', level=2)

doc.add_paragraph(
    '상품 목록 페이지에서 카테고리·정렬·검색 값을 useState와 URL searchParams 두 곳에서 '
    '동시에 관리했습니다. Header에서 카테고리 링크를 클릭하면 URL은 바뀌지만 useState는 '
    '그대로여서 필터가 작동하지 않는 버그가 발생했습니다.'
)
doc.add_paragraph('URL에 반영되는 데이터는 useState로 복제하지 말 것 → searchParams.get()으로 직접 읽기', style='List Bullet')
doc.add_paragraph('값 변경은 router.push()로 URL을 업데이트 → 컴포넌트가 자동 리렌더링', style='List Bullet')
doc.add_paragraph('같은 데이터를 두 곳에서 관리하면 반드시 동기화 버그가 발생함', style='List Bullet')

# 4-3. next/image 선행 설계
doc.add_heading('4-3. 이미지 컴포넌트 선행 설계', level=2)

doc.add_paragraph(
    '개발 초기에 실제 이미지 파일 없이 진행하면서, 이미지 깨짐 문제가 반복되었습니다. '
    'next/image에 onError 폴백 구조를 처음부터 설계하지 않았기 때문입니다.'
)
doc.add_paragraph('이미지 파일이 없어도 처음부터 next/image + onError 폴백 구조로 작성할 것', style='List Bullet')
doc.add_paragraph('placehold.co를 폴백 URL로 사용하면 개발 중에도 깨지지 않는 UI 유지 가능', style='List Bullet')
doc.add_paragraph('이미지는 나중에 넣되, 넣는 순간 바로 동작하는 구조를 처음부터 갖출 것', style='List Bullet')

# 4-4. Next.js App Router 주의사항
doc.add_heading('4-4. Next.js App Router 주의사항', level=2)

doc.add_paragraph(
    'Next.js 14 App Router는 서버 컴포넌트가 기본입니다. useState, useEffect, useSearchParams 등 '
    '클라이언트 기능을 사용하려면 반드시 \'use client\'를 선언해야 하고, useSearchParams는 '
    'Suspense 바운더리 안에서 사용해야 빌드 에러가 발생하지 않습니다.'
)
doc.add_paragraph("'use client'는 상태/이벤트/localStorage를 사용하는 컴포넌트에만 최소한으로 붙이기", style='List Bullet')
doc.add_paragraph('useSearchParams 사용 시 반드시 Suspense 래퍼로 감싸기', style='List Bullet')
doc.add_paragraph('next/font가 빌드 환경에서 실패할 수 있음 → CSS @import 방식을 대안으로 준비', style='List Bullet')

# 4-5. Tailwind 커스텀 테마 전략
doc.add_heading('4-5. Tailwind 커스텀 테마 전략', level=2)

doc.add_paragraph(
    'tailwind.config.js에 커스텀 색상(cream, chocolate, caramel, gold, rose), '
    '커스텀 폰트(display, body, accent), 커스텀 그림자(warm-sm~xl)를 프로젝트 시작 시 '
    '정의했습니다. 덕분에 개발 전 과정에서 하드코딩 색상 0건, 디자인 일관성을 유지할 수 있었습니다.'
)
doc.add_paragraph('색상·폰트·그림자를 프로젝트 1단계에서 tailwind.config.js에 확정할 것', style='List Bullet')
doc.add_paragraph('Tailwind 기본 색상(red-500, blue-600 등) 대신 시맨틱 이름(bg-cream, text-chocolate) 사용', style='List Bullet')
doc.add_paragraph('다크모드는 dark: 변형을 활용하되, dm-bg, dm-surface 등 다크 전용 색상도 미리 정의할 것', style='List Bullet')

# 4-6. 다크모드 사후 추가의 부담
doc.add_heading('4-6. 다크모드 사후 추가의 부담', level=2)

doc.add_paragraph(
    '라이트 모드 개발을 완료한 후 다크모드를 추가했습니다. 13개 페이지와 23개 컴포넌트에 '
    'dark: 클래스를 일괄 추가해야 했는데, 2차 프로젝트에서 CSS 하드코딩을 사후에 변환한 것과 '
    '유사한 패턴이었습니다. Tailwind 덕분에 CSS 변수 전환보다 훨씬 간단했지만, '
    '처음부터 dark: 클래스를 함께 작성했다면 더 효율적이었을 것입니다.'
)
doc.add_paragraph('다크모드 지원 계획이 있다면 컴포넌트 작성 시 dark: 클래스를 처음부터 함께 넣을 것', style='List Bullet')
doc.add_paragraph('tailwind.config.js에 다크모드 색상(dm-bg, dm-surface, dm-card 등)을 사전 정의한 것은 정답이었음', style='List Bullet')

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 5. 결론
# ═══════════════════════════════════════════════════════

doc.add_heading('5. 결론', level=1)

doc.add_paragraph(
    'agents_homepage는 1차(rndpark), 2차(module_homepage)의 교훈을 적극 반영한 세 번째 프로젝트입니다. '
    '에뮬레이터-프로덕션 차이(0건), 배포 설정 실패(0건), Auth 재인증(0건), CSS 하드코딩(0건), '
    '모듈 연동 충돌(0건), 변수 중복(0건) 등 6가지를 완전히 예방했습니다.'
)

doc.add_paragraph(
    '반면, 스타일 사후 수정(다크모드 일괄 추가)과 컨텍스트 소진은 패턴이 반복되었으나 '
    '영향도가 크게 감소했습니다. 특히 Tailwind CSS 덕분에 다크모드 추가가 CSS 변수 전환보다 '
    '훨씬 간단했고, Claude Code 에이전트 덕분에 컨텍스트 소진 빈도가 대폭 줄었습니다.'
)

doc.add_paragraph(
    '또한 URL 파라미터 단일 소스 원칙, 이미지 선행 설계, Next.js Suspense 처리, '
    'Tailwind 커스텀 테마 전략 등 이번 프로젝트에서 새로 발견한 교훈도 있었습니다. '
    '이 경험들을 CLAUDE.md에 규칙으로 반영하여, 다음 프로젝트에서도 재활용할 수 있도록 정리했습니다.'
)

doc.add_paragraph(
    '13개 페이지와 23개 컴포넌트를 모두 완성하고, 4개 Context로 전역 상태 관리를 구현했으며, '
    '다크모드, 접근성(ARIA, 키보드 네비게이션), SEO 메타데이터, 성능 최적화(dynamic import, next/image)까지 '
    '적용하여 Vercel 자동 배포를 완료했습니다.'
)


# ═══════════════════════════════════════════════════════
# 파일 저장
# ═══════════════════════════════════════════════════════

output_path = '/home/user/agents_homepage/에이전트홈페이지_통합회고보고서.docx'
doc.save(output_path)
print(f'보고서 생성 완료: {output_path}')
